"""
Tạo file HDSD_CapNhatTuan_AI_USSPTD.docx
Chạy: python scripts/gen_hdsd_capnhattuan.py
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

BASE   = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
EVD    = os.path.join(BASE, 'evd', 'weekly-update')
OUT    = os.path.join(BASE, 'HDSD_CapNhatTuan_AI_USSPTD.docx')

# ── Màu sắc ──────────────────────────────────────────────────────────────────
C_PRIMARY   = RGBColor(0x7B, 0x2C, 0xBF)   # tím chủ đạo
C_DARK      = RGBColor(0x1A, 0x1A, 0x2E)   # tiêu đề đậm
C_GRAY      = RGBColor(0x6B, 0x72, 0x80)   # mô tả phụ
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_NOTE_BG   = RGBColor(0xF5, 0xF0, 0xFF)   # nền chú thích tím nhạt
C_WARN_BG   = RGBColor(0xFF, 0xF3, 0xE0)   # nền cảnh báo cam
C_OK_BG     = RGBColor(0xE8, 0xF5, 0xE9)   # nền thành công xanh

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_para_bg(para, hex_color):
    """Đặt shading cho đoạn văn."""
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)

def add_run(para, text, bold=False, italic=False, color=None, size=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color: run.font.color.rgb = color
    if size:  run.font.size = Pt(size)
    return run

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(16)
    run.font.color.rgb = C_PRIMARY
    # Đường kẻ dưới
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),  'single')
    bot.set(qn('w:sz'),   '6')
    bot.set(qn('w:space'),'4')
    bot.set(qn('w:color'),'7B2CBF')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(13)
    run.font.color.rgb = C_DARK
    return p

def body(doc, text, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = C_DARK
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.6)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = C_DARK
    return p

def note_box(doc, title, text, bg='F5F0FF', title_color='7B2CBF'):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.width = Cm(16)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    add_run(p1, title + '  ', bold=True, color=RGBColor.from_string(title_color), size=10)
    p2 = cell.add_paragraph()
    add_run(p2, text, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_image(doc, filename, width_cm=15, caption=None):
    path = os.path.join(EVD, filename)
    if not os.path.exists(path):
        body(doc, f'[Ảnh: {filename} — không tìm thấy]')
        return
    doc.add_picture(path, width=Cm(width_cm))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        for run in cp.runs:
            run.font.size   = Pt(9)
            run.font.italic = True
            run.font.color.rgb = C_GRAY

def page_break(doc):
    doc.add_page_break()

# ── Bắt đầu tạo document ─────────────────────────────────────────────────────

doc = Document()

# Margin A4
sec = doc.sections[0]
sec.page_width  = Cm(21)
sec.page_height = Cm(29.7)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.5)
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.5)

# Default font
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# TRANG BÌA
# ══════════════════════════════════════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_title.add_run('HƯỚNG DẪN SỬ DỤNG')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = C_PRIMARY

cover_sub = doc.add_paragraph()
cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = cover_sub.add_run('Tính năng CẬP NHẬT TIẾN ĐỘ TUẦN')
run2.bold = True
run2.font.size = Pt(18)
run2.font.color.rgb = C_DARK

doc.add_paragraph()

cover_system = doc.add_paragraph()
cover_system.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(cover_system, 'Hệ thống Bình dân hóa AI — AI Use Case Platform (AI-US-SPTD)', color=C_GRAY, size=12)

for _ in range(6):
    doc.add_paragraph()

meta_tbl = doc.add_table(rows=4, cols=2)
meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('Phiên bản',    'v1.0 — 2026'),
    ('Ngày tạo',     datetime.date.today().strftime('%d/%m/%Y')),
    ('Đối tượng',    'Tất cả nhân viên sử dụng hệ thống AI-US-SPTD'),
    ('Liên hệ hỗ trợ', 'Ban Chuyển đổi số / Champion AI'),
]
for i, (k, v) in enumerate(meta_data):
    row = meta_tbl.rows[i]
    row.cells[0].paragraphs[0].add_run(k).bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    row.cells[1].paragraphs[0].add_run(v).font.size = Pt(11)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 1. TỔNG QUAN
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, '1. Tổng quan tính năng')

body(doc,
    'Tính năng Cập nhật tiến độ tuần cho phép chủ sở hữu Use Case báo cáo '
    'tiến độ thực hiện mỗi tuần, cập nhật số liệu thực tế và đề xuất nâng '
    'Stage khi đủ điều kiện. Hệ thống tự động tính điểm và lưu lịch sử mỗi '
    'lần cập nhật.')

heading2(doc, 'Vai trò và quyền hạn')

role_tbl = doc.add_table(rows=4, cols=3)
role_tbl.style = 'Table Grid'
role_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

headers = ['Vai trò', 'Phạm vi Use Case có thể cập nhật', 'Ghi chú']
for i, h in enumerate(headers):
    cell = role_tbl.rows[0].cells[i]
    set_cell_bg(cell, '7B2CBF')
    p = cell.paragraphs[0]
    add_run(p, h, bold=True, color=C_WHITE, size=10.5)

rows_data = [
    ('Người dùng (User)',    'Chỉ các UC do chính mình là Chủ sở hữu',             'Xem và cập nhật UC của bản thân'),
    ('Champion',             'Tất cả UC trong cùng Team',                           'Hỗ trợ theo dõi tiến độ toàn team'),
    ('Admin',                'Toàn bộ UC trong hệ thống (117+ UC)',                 'Không giới hạn phạm vi'),
]
for i, (role, scope, note) in enumerate(rows_data, 1):
    bg = 'F5F0FF' if i % 2 == 0 else 'FFFFFF'
    cells = role_tbl.rows[i].cells
    set_cell_bg(cells[0], bg); cells[0].paragraphs[0].add_run(role).font.size = Pt(10.5)
    set_cell_bg(cells[1], bg); cells[1].paragraphs[0].add_run(scope).font.size = Pt(10.5)
    set_cell_bg(cells[2], bg); cells[2].paragraphs[0].add_run(note).font.size = Pt(10.5)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

heading2(doc, 'Vòng đời Stage (giai đoạn)')

stage_tbl = doc.add_table(rows=5, cols=3)
stage_tbl.style = 'Table Grid'
stage_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

sh = ['Stage', 'Tên', 'Ý nghĩa']
for i, h in enumerate(sh):
    cell = stage_tbl.rows[0].cells[i]
    set_cell_bg(cell, '1A1A2E')
    add_run(cell.paragraphs[0], h, bold=True, color=C_WHITE, size=10.5)

stage_data = [
    ('S1', 'Ý tưởng',   'UC mới được đăng ký, chưa có kết quả thực tế'),
    ('S2', 'Pilot',     'Đang thử nghiệm với nhóm nhỏ, thu thập feedback'),
    ('S3', 'Chuẩn hóa', 'Đã có quy trình, mở rộng trong bộ phận'),
    ('S4', 'Scale',     'Giai đoạn cao nhất — nhân rộng đa bộ phận / toàn tổ chức'),
]
stage_colors = ['E8F4FD', 'E8F5E9', 'FFF8E1', 'F3E5F5']
for i, (code, name, desc) in enumerate(stage_data, 1):
    cells = stage_tbl.rows[i].cells
    set_cell_bg(cells[0], stage_colors[i-1])
    add_run(cells[0].paragraphs[0], code, bold=True, size=11)
    cells[1].paragraphs[0].add_run(name).font.size = Pt(10.5)
    cells[2].paragraphs[0].add_run(desc).font.size = Pt(10.5)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 2. TRUY CẬP
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, '2. Truy cập tính năng')

body(doc,
    'Từ bất kỳ trang nào trong hệ thống, nhấn vào mục "Cập nhật tuần" '
    'trong thanh menu bên trái (phần AI GOVERNANCE).')

add_image(doc, 'T01-page-load.png', width_cm=14,
          caption='Hình 1 — Giao diện trang Cập nhật tiến độ tuần khi mới mở')

note_box(doc,
    '📌  Lưu ý:',
    'Trang chỉ hiển thị ô chọn Use Case khi mới vào. '
    'Form nhập liệu sẽ xuất hiện SAU KHI bạn đã chọn một Use Case.',
    bg='FFF8E1', title_color='E65100')

# ══════════════════════════════════════════════════════════════════════════════
# 3. CHỌN USE CASE
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, '3. Bước 1 — Chọn Use Case cần cập nhật')

body(doc,
    'Nhấn vào ô "— Nhấn để chọn use case —" để mở cửa sổ chọn Use Case. '
    'Hệ thống sẽ hiển thị danh sách các Use Case bạn có quyền cập nhật '
    'theo vai trò của mình.')

heading2(doc, '3.1  Cửa sổ chọn Use Case')

add_image(doc, 'T02-picker-modal-open.png', width_cm=14,
          caption='Hình 2 — Cửa sổ chọn Use Case dạng bảng với đầy đủ thông tin')

body(doc, 'Bảng hiển thị các cột:')
for col in ['Mã UC (AIUS-XXXX)', 'Tên Use Case và trạng thái phê duyệt',
            'Chủ sở hữu', 'Team', 'Stage hiện tại', 'Tiến độ (%)', 'Cập nhật lần cuối']:
    bullet(doc, col)

doc.add_paragraph().paragraph_format.space_after = Pt(4)
p = doc.add_paragraph()
add_run(p, '💡  Góc trên phải cửa sổ', bold=True, color=C_PRIMARY, size=10.5)
add_run(p, ' hiển thị phạm vi của bạn — ví dụ: ', size=10.5)
add_run(p, '"Admin · Toàn bộ"', bold=True, size=10.5)
add_run(p, ' hoặc ', size=10.5)
add_run(p, '"Champion · Team Số"', bold=True, size=10.5)
add_run(p, '.', size=10.5)
p.paragraph_format.space_after = Pt(8)

heading2(doc, '3.2  Tìm kiếm và lọc')

add_image(doc, 'T03-picker-filter.png', width_cm=14,
          caption='Hình 3 — Tìm kiếm theo từ khóa (tên UC, mã AIUS, chủ sở hữu)')

body(doc, 'Có 3 cách lọc độc lập, có thể kết hợp:')

filter_tbl = doc.add_table(rows=4, cols=2)
filter_tbl.style = 'Table Grid'
fh = ['Bộ lọc', 'Mô tả']
for i, h in enumerate(fh):
    cell = filter_tbl.rows[0].cells[i]
    set_cell_bg(cell, '7B2CBF')
    add_run(cell.paragraphs[0], h, bold=True, color=C_WHITE, size=10.5)

filter_data = [
    ('Ô tìm kiếm (🔍)',    'Nhập tên UC, mã AIUS-XXXX hoặc tên chủ sở hữu — kết quả lọc ngay tức thì'),
    ('Lọc theo Stage',     'Chọn S1 / S2 / S3 / S4 hoặc "Tất cả Stage" để hiển thị nhóm cần quan tâm'),
    ('Lọc trạng thái CN',  'Chọn "Hôm nay" / "Tuần này" / "Chưa có" để tìm nhanh UC cần cập nhật gấp'),
]
for i, (f, d) in enumerate(filter_data, 1):
    bg = 'F5F0FF' if i % 2 == 0 else 'FFFFFF'
    cells = filter_tbl.rows[i].cells
    set_cell_bg(cells[0], bg); add_run(cells[0].paragraphs[0], f, bold=True, size=10.5)
    set_cell_bg(cells[1], bg); cells[1].paragraphs[0].add_run(d).font.size = Pt(10.5)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_image(doc, 'T04-stage-filter.png', width_cm=14,
          caption='Hình 4 — Lọc theo Stage (ví dụ: S3 Chuẩn hóa, S4 Scale)')

heading2(doc, '3.3  Chọn một Use Case')

body(doc,
    'Nhấn vào bất kỳ dòng nào trong bảng để chọn Use Case đó. '
    'Cửa sổ sẽ tự đóng và tên UC hiện ngay trên thanh chọn. '
    'Nhấn ✕ để bỏ chọn và chọn lại Use Case khác.')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 4. XEM THÔNG TIN UC
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, '4. Bước 2 — Xem thông tin Use Case đã chọn')

add_image(doc, 'T05-uc-selected.png', width_cm=14,
          caption='Hình 5 — Sau khi chọn UC: thẻ thông tin, phần Stage, và form nhập liệu hiển thị')

body(doc, 'Sau khi chọn, hệ thống hiển thị:')
bullet(doc, 'Thẻ thông tin UC: tên đầy đủ, mã AIUS, team, trạng thái phê duyệt')
bullet(doc, 'Điểm hiện tại (xx/100) và Stage hiện tại')
bullet(doc, 'Thời điểm cập nhật lần cuối — nếu quá 7 ngày sẽ có cảnh báo màu cam')
bullet(doc, 'Phần Stage (giai đoạn) và Form nhập liệu bên dưới')

note_box(doc,
    '⚠️  Cảnh báo quá hạn:',
    'Nếu Use Case chưa được cập nhật quá 7 ngày, hệ thống hiển thị banner '
    'màu đỏ nhắc nhở. Cập nhật thường xuyên giúp duy trì điểm Governance.',
    bg='FFF3E0', title_color='E65100')

# ══════════════════════════════════════════════════════════════════════════════
# 5. ĐIỀN TIẾN ĐỘ
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, '5. Bước 3 — Điền thông tin tiến độ tuần')

body(doc,
    'Form gồm 3 nhóm trường chính. Các trường đánh dấu * là bắt buộc.')

heading2(doc, '5.1  Tiến độ thực hiện')

add_image(doc, 'T08-progress-slider.png', width_cm=14,
          caption='Hình 6 — Thanh kéo % hoàn thành cập nhật ngay tức thì')

form_tbl = doc.add_table(rows=6, cols=3)
form_tbl.style = 'Table Grid'
form_headers = ['Trường', 'Mô tả', 'Bắt buộc']
for i, h in enumerate(form_headers):
    cell = form_tbl.rows[0].cells[i]
    set_cell_bg(cell, '1A1A2E')
    add_run(cell.paragraphs[0], h, bold=True, color=C_WHITE, size=10.5)

fields = [
    ('% Hoàn thành',               'Kéo thanh trượt từ 0–100%. Hiển thị ngay lập tức.',                      '✅'),
    ('Cập nhật tuần này *',        'Mô tả công việc thực hiện trong tuần: kết quả, số liệu, điểm nổi bật.',  '✅'),
    ('Milestone tiếp theo',        'Mục tiêu cụ thể cho tuần tới hoặc cột mốc quan trọng.',                  '—'),
    ('Số lần sử dụng trong tháng', 'Số lượt thực tế AI được áp dụng trong tháng (số nguyên).',               '—'),
    ('Giờ tiết kiệm thực tế/tháng','Ước tính giờ tiết kiệm được nhờ AI (số nguyên, đơn vị: giờ/tháng).',    '—'),
]
for i, (f, d, req) in enumerate(fields, 1):
    bg = 'F5F0FF' if i % 2 == 0 else 'FFFFFF'
    cells = form_tbl.rows[i].cells
    set_cell_bg(cells[0], bg); add_run(cells[0].paragraphs[0], f, bold=(req=='✅'), size=10.5)
    set_cell_bg(cells[1], bg); cells[1].paragraphs[0].add_run(d).font.size = Pt(10.5)
    set_cell_bg(cells[2], bg)
    r = cells[2].paragraphs[0].add_run(req)
    r.font.size = Pt(11)
    if req == '✅': r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

heading2(doc, '5.2  Rào cản & Hỗ trợ')

body(doc,
    'Điền vào 2 ô tùy chọn: Rào cản hiện tại (khó khăn về kỹ thuật, quy trình, '
    'con người...) và Cần hỗ trợ từ quản lý. Thông tin này giúp Champion và '
    'quản lý nắm bắt nhanh để can thiệp kịp thời.')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 6. ĐỀ XUẤT NÂNG STAGE
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, '6. Bước 4 — Đề xuất nâng Stage (tùy chọn)')

body(doc,
    'Khi Use Case đủ điều kiện chuyển lên Stage tiếp theo, bạn có thể đề xuất '
    'nâng Stage ngay trong lần cập nhật tuần. Tích vào checkbox '
    '"Đề xuất nâng stage lần này" để mở bảng điều kiện.')

heading2(doc, '6.1  Bật chế độ đề xuất nâng Stage')

add_image(doc, 'T05-stage-upgrade-panel.png', width_cm=14,
          caption='Hình 7 — Panel đề xuất nâng Stage với Stage đích và checklist điều kiện')

body(doc,
    'Hệ thống tự động hiển thị Stage đích (Stage liền kề tiếp theo) '
    'và danh sách 4 điều kiện cần xác nhận trước khi được phép nâng.')

heading2(doc, '6.2  Checklist điều kiện')

add_image(doc, 'T09-checklist-validation.png', width_cm=14,
          caption='Hình 8 — Cảnh báo khi chưa tích đủ checklist điều kiện nâng Stage')

body(doc,
    'Mỗi Stage có 4 tiêu chí riêng. Bạn phải tích đủ cả 4 mới có thể gửi. '
    'Nếu nhấn "Gửi cập nhật" khi còn bỏ sót, hệ thống hiển thị cảnh báo màu đỏ.')

checklist_tbl = doc.add_table(rows=4, cols=3)
checklist_tbl.style = 'Table Grid'
ch = ['Stage đích', 'Điều kiện cần xác nhận (4/4)', 'Chú thích']
for i, h in enumerate(ch):
    cell = checklist_tbl.rows[0].cells[i]
    set_cell_bg(cell, '7B2CBF')
    add_run(cell.paragraphs[0], h, bold=True, color=C_WHITE, size=10)

checklists = [
    ('S2 · Pilot',
     '1. Đã pilot với ít nhất 3 người dùng\n2. Có feedback thực tế từ người dùng\n3. Tỷ lệ hoàn thành ≥ 50%\n4. Có kết quả đo lường ban đầu',
     'Từ S1 → S2'),
    ('S3 · Chuẩn hóa',
     '1. Đã pilot với ít nhất 5 người dùng\n2. Monthly Usage ≥ 10 lần/tháng\n3. Có SOP / tài liệu hướng dẫn\n4. Tỷ lệ hoàn thành ≥ 70%',
     'Từ S2 → S3'),
    ('S4 · Scale',
     '1. Đã phổ biến với ít nhất 5 người dùng trong team\n2. Monthly Usage ≥ 30 lần/tháng\n3. Có feedback đối tác/đơn vị liên quan\n4. Tỷ lệ hoàn thành ≥ 80%',
     'Từ S3 → S4'),
]
for i, (stage, clist, note) in enumerate(checklists, 1):
    bg = 'F5F0FF' if i % 2 == 0 else 'FFFFFF'
    cells = checklist_tbl.rows[i].cells
    set_cell_bg(cells[0], bg); add_run(cells[0].paragraphs[0], stage, bold=True, size=10.5)
    set_cell_bg(cells[1], bg); cells[1].paragraphs[0].add_run(clist).font.size = Pt(10)
    set_cell_bg(cells[2], bg); cells[2].paragraphs[0].add_run(note).font.size = Pt(10.5)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

heading2(doc, '6.3  Trường bổ sung khi nâng lên S4 · Scale')

add_image(doc, 'T07-s4-fields.png', width_cm=14,
          caption='Hình 9 — Các trường bổ sung bắt buộc khi đề xuất nâng lên S4 · Scale')

body(doc,
    'Khi đề xuất nâng lên S4, hệ thống yêu cầu thêm 2 trường quan trọng:')
bullet(doc, 'Kế hoạch mở rộng (Scale Plan): mô tả cụ thể việc nhân rộng sang đơn vị khác, '
            'mục tiêu số lượng người dùng, timeline.')
bullet(doc, 'Rủi ro khi scale: các rào cản về bảo mật, tuân thủ, nhân lực cần giải quyết trước khi mở rộng.')

note_box(doc,
    '📋  Lưu ý S4:',
    'Các trường S4 chỉ được lưu vào WEEKLY_LOG (lịch sử) để tham khảo, '
    'không ghi đè dữ liệu gốc của Use Case. Stage chính thức vẫn được cập nhật vào MASTER_DATA.',
    bg='F5F0FF', title_color='7B2CBF')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 7. GỬI CẬP NHẬT
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, '7. Bước 5 — Gửi cập nhật')

body(doc,
    'Sau khi điền đầy đủ thông tin, nhấn nút "Gửi cập nhật" ở cuối trang. '
    'Hệ thống sẽ xử lý và hiển thị kết quả trong vài giây.')

heading2(doc, '7.1  Cập nhật thành công (không nâng Stage)')

add_image(doc, 'T10-submit-success.png', width_cm=14,
          caption='Hình 10 — Thông báo cập nhật thành công, hiển thị điểm mới')

body(doc,
    'Sau khi gửi thành công, hệ thống hiển thị:\n'
    '• Điểm mới (ví dụ: Điểm mới: 57/100)\n'
    '• Dòng lịch sử mới nhất xuất hiện trong Timeline phía dưới')

heading2(doc, '7.2  Cập nhật thành công có nâng Stage')

add_image(doc, 'T11-stage-submit-success.png', width_cm=14,
          caption='Hình 11 — Thông báo thành công kèm Stage mới khi được nâng (ví dụ: S4 · Scale)')

body(doc,
    'Khi đề xuất nâng Stage thành công, thông báo hiển thị thêm:\n'
    '• Stage mới (ví dụ: Stage mới: S4 · Scale)\n'
    '• Thẻ UC card tự động cập nhật Stage mới\n'
    '• Timeline ghi lại sự kiện nâng Stage với màu tím đặc biệt')

note_box(doc,
    '✅  Điểm được tính tự động:',
    'Hệ thống tự động tính lại điểm Governance sau mỗi lần cập nhật, '
    'bao gồm điểm Cập nhật thường xuyên, điểm Số liệu thực tế, điểm Stage. '
    'Điểm này ảnh hưởng trực tiếp đến thứ hạng trên Leaderboard.',
    bg='E8F5E9', title_color='2E7D32')

# ══════════════════════════════════════════════════════════════════════════════
# 8. LỊCH SỬ CẬP NHẬT (TIMELINE)
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, '8. Xem lịch sử cập nhật (Timeline)')

add_image(doc, 'T09-timeline.png', width_cm=14,
          caption='Hình 12 — Timeline lịch sử cập nhật của Use Case, sắp xếp mới nhất lên đầu')

body(doc,
    'Timeline hiển thị toàn bộ lịch sử cập nhật của Use Case được chọn, '
    'sắp xếp theo thứ tự mới nhất lên đầu. Mỗi thẻ trong Timeline gồm:')
bullet(doc, 'Ngày giờ cập nhật')
bullet(doc, 'Phần trăm hoàn thành tại thời điểm đó')
bullet(doc, 'Nội dung cập nhật tuần')
bullet(doc, 'Số liệu thực tế (nếu có): lượt sử dụng, giờ tiết kiệm, lần tái sử dụng')
bullet(doc, 'Milestone tiếp theo (nếu có)')

body(doc,
    'Các thẻ có sự kiện nâng Stage được tô màu tím và có nhãn chuyển Stage '
    '(ví dụ: S3 → S4) để dễ theo dõi vòng đời UC.', space_after=8)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 9. THAO TÁC KHÁC
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, '9. Các thao tác khác')

heading2(doc, '9.1  Bỏ chọn Use Case và chọn lại')

add_image(doc, 'T06-clear-selection.png', width_cm=12,
          caption='Hình 13 — Nhấn ✕ để bỏ chọn Use Case và quay về màn hình ban đầu')

body(doc,
    'Nhấn nút ✕ nhỏ ngay trên thanh chọn để bỏ chọn Use Case hiện tại. '
    'Form và thông tin sẽ được xóa sạch, bạn có thể mở lại cửa sổ chọn '
    'Use Case để chọn UC khác.')

heading2(doc, '9.2  Đóng cửa sổ chọn mà không chọn')

add_image(doc, 'T07-modal-close-escape.png', width_cm=12,
          caption='Hình 14 — Nhấn phím Escape hoặc nút ✕ để đóng cửa sổ không cần chọn')

body(doc,
    'Để đóng cửa sổ chọn Use Case mà không thay đổi lựa chọn hiện tại:')
bullet(doc, 'Nhấn phím Escape trên bàn phím')
bullet(doc, 'Nhấn nút ✕ ở góc phải cửa sổ')
bullet(doc, 'Nhấn vào vùng mờ bên ngoài cửa sổ')

heading2(doc, '9.3  Làm lại (reset)')

body(doc,
    'Nhấn nút "Làm lại" ở cuối form để xóa toàn bộ dữ liệu đã nhập '
    'và trở về trạng thái ban đầu (không bỏ chọn Use Case). '
    'Hữu ích khi muốn nhập lại từ đầu.')

# ══════════════════════════════════════════════════════════════════════════════
# 10. LƯU Ý VÀ FAQ
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, '10. Lưu ý quan trọng & FAQ')

faq_data = [
    ('Tôi không thấy UC của mình trong danh sách?',
     'Kiểm tra lại: (1) Tên chủ sở hữu trong hệ thống có khớp với email đăng nhập không. '
     '(2) Thử xóa bộ lọc Stage/trạng thái và tìm kiếm lại. '
     '(3) Liên hệ Admin để kiểm tra thông tin tài khoản.'),
    ('Có thể cập nhật nhiều lần trong một tuần không?',
     'Có. Hệ thống không giới hạn số lần cập nhật. Mỗi lần gửi đều được lưu vào Timeline. '
     'Tuy nhiên điểm "Cập nhật thường xuyên" chỉ tính 1 lần/tuần.'),
    ('Điểm thay đổi như thế nào sau khi cập nhật?',
     'Hệ thống tự tính lại điểm ngay sau khi gửi. Điểm bao gồm: Governance score '
     '(cập nhật đều), Impact score (số liệu thực tế), Stage score. '
     'Kết quả hiển thị ngay trên màn hình thành công.'),
    ('Tôi đề xuất nâng Stage nhưng chưa đủ điều kiện thực tế?',
     'Hệ thống dựa vào sự xác nhận tự giác của bạn qua checklist. '
     'Champion và Admin có thể kiểm tra và yêu cầu giải trình. '
     'Hãy chỉ đề xuất khi thực sự đã đáp ứng đủ tiêu chí.'),
    ('Tôi có thể xem lịch sử của UC người khác không?',
     'Không. Timeline chỉ hiển thị lịch sử của UC bạn đang chọn. '
     'Bạn chỉ xem được UC trong phạm vi quyền hạn (xem mục Vai trò & Quyền hạn).'),
    ('Sau khi nâng lên S4, có thể cập nhật tuần tiếp không?',
     'Có. S4 là Stage cao nhất nhưng vẫn cần duy trì cập nhật hàng tuần. '
     'Ô đề xuất nâng Stage sẽ ẩn đi và thay bằng thông báo "Use case đã đạt S4 — Stage cao nhất".'),
]

for q, a in faq_data:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, 'Q: ' + q, bold=True, color=C_PRIMARY, size=11)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.left_indent = Cm(0.5)
    add_run(p2, 'A: ' + a, size=11)

# ══════════════════════════════════════════════════════════════════════════════
# 11. QUY TRÌNH TÓM TẮT
# ══════════════════════════════════════════════════════════════════════════════

page_break(doc)

heading1(doc, '11. Quy trình tóm tắt')

steps = [
    ('Bước 1', 'Vào "Cập nhật tuần" từ menu bên trái'),
    ('Bước 2', 'Nhấn ô chọn → Tìm và chọn Use Case trong bảng'),
    ('Bước 3', 'Kiểm tra thông tin UC (Stage, điểm, lần cập nhật cuối)'),
    ('Bước 4', 'Điền tiến độ: % hoàn thành, nội dung tuần này, số liệu thực tế'),
    ('Bước 5', '[Tùy chọn] Tích "Đề xuất nâng Stage" → Xác nhận 4 điều kiện'),
    ('Bước 6', 'Nhấn "Gửi cập nhật" → Hệ thống tính điểm và lưu vào Timeline'),
]

flow_tbl = doc.add_table(rows=len(steps), cols=2)
flow_tbl.style = 'Table Grid'
flow_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (step, desc) in enumerate(steps):
    cells = flow_tbl.rows[i].cells
    set_cell_bg(cells[0], '7B2CBF' if i == 0 else ('5B21B6' if i == len(steps)-1 else 'F5F0FF'))
    p0 = cells[0].paragraphs[0]
    r0 = add_run(p0, step, bold=True, size=11,
                 color=C_WHITE if i in (0, len(steps)-1) else C_PRIMARY)
    set_cell_bg(cells[1], 'FFFFFF')
    cells[1].paragraphs[0].add_run(desc).font.size = Pt(11)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

note_box(doc,
    '📅  Tần suất khuyến nghị:',
    'Cập nhật mỗi tuần vào cuối tuần (Thứ Sáu hoặc Thứ Bảy) để đảm bảo '
    'điểm Governance được tính đầy đủ và Champion nắm bắt tiến độ kịp thời.',
    bg='E8F5E9', title_color='2E7D32')

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_foot,
        f'Tài liệu này được tạo tự động — {datetime.date.today().strftime("%d/%m/%Y")}  |  '
        'Hệ thống AI Use Case Platform v3.10 — TT SPTD',
        color=C_GRAY, size=9, italic=True)

# ── Lưu file ─────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f'OK  Da tao: {OUT}')
