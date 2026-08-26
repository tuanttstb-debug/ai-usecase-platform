# -*- coding: utf-8 -*-
"""
Dựng HƯỚNG DẪN SỬ DỤNG nền tảng Bình dân hóa AI H2/2026 — cho TEAMLEAD + NHÂN SỰ.
Ảnh minh họa: screenshots/h2/*.png (chụp qua tests/zz-capture-h2-guide.spec.js).
Output: HDSD_H2_2026_Teamlead_NhanSu.docx
"""
import os, datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(__file__)
SHOTS = os.path.join(BASE, 'screenshots', 'h2')
OUT = os.path.join(BASE, 'HDSD_H2_2026_Teamlead_NhanSu.docx')

PURPLE = RGBColor(0x6B, 0x21, 0xA8)
DARK   = RGBColor(0x1F, 0x25, 0x37)
GRAY   = RGBColor(0x6B, 0x72, 0x80)
GREEN  = RGBColor(0x05, 0x96, 0x69)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=PURPLE):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
    return p


def add_body(doc, text, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.bold = bold
    if color: run.font.color.rgb = color
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(it); r.font.size = Pt(10.5)


def add_image(doc, filename, caption, width=Cm(15.5)):
    path = os.path.join(SHOTS, filename)
    if not os.path.exists(path):
        add_body(doc, f'[Hình: {caption} — thiếu file {filename}]', color=GRAY, size=9)
        return
    doc.add_picture(path, width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.runs[0] if cap.runs else cap.add_run(caption)
    r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = GRAY


def add_step_box(doc, num, title, desc):
    tbl = doc.add_table(rows=1, cols=2); tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c0 = tbl.rows[0].cells[0]; c0.width = Cm(1.3); set_cell_bg(c0, '6B21A8')
    p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(str(num)); r0.font.size = Pt(15); r0.font.bold = True; r0.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    c1 = tbl.rows[0].cells[1]; p1 = c1.paragraphs[0]
    ra = p1.add_run(title + '\n'); ra.font.bold = True; ra.font.size = Pt(11)
    rb = p1.add_run(desc); rb.font.size = Pt(10); rb.font.color.rgb = GRAY
    doc.add_paragraph()


def add_note_box(doc, text, kind='info'):
    cmap = {'info': 'EDE9FE', 'warn': 'FEF3C7', 'tip': 'D1FAE5'}
    tbl = doc.add_table(rows=1, cols=1); tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]; set_cell_bg(cell, cmap.get(kind, 'EDE9FE'))
    r = cell.paragraphs[0].add_run(text); r.font.size = Pt(10); r.font.color.rgb = DARK
    doc.add_paragraph()


def add_table(doc, headers, rows, col_fill='6B21A8'):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers)); tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]; set_cell_bg(cell, col_fill)
        r = cell.paragraphs[0].add_run(h); r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for ri, rd in enumerate(rows, start=1):
        for ci, val in enumerate(rd):
            r = tbl.rows[ri].cells[ci].paragraphs[0].add_run(str(val)); r.font.size = Pt(9.5)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles['Normal']; style.font.name = 'Segoe UI'; style.font.size = Pt(11)

# ── Trang bìa ──
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('HƯỚNG DẪN SỬ DỤNG'); r.font.size = Pt(24); r.font.bold = True; r.font.color.rgb = PURPLE
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('Nền tảng Bình dân hóa AI — Chương trình H2/2026'); r.font.size = Pt(14); r.font.color.rgb = DARK
s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s2.add_run('Dành cho Teamlead & Nhân sự · TT SPTD'); r.font.size = Pt(12); r.font.color.rgb = GRAY
d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = d.add_run('Cập nhật: ' + datetime.date.today().strftime('%d/%m/%Y')); r.font.size = Pt(10); r.font.italic = True; r.font.color.rgb = GRAY
doc.add_paragraph()

# ── 1. Giới thiệu ──
add_heading(doc, '1. Giới thiệu', 1)
add_body(doc, 'Nền tảng ghi nhận và quản trị việc ứng dụng AI trong công việc theo chương trình '
              '"Bình dân hóa AI" nửa cuối 2026 (H2). Mỗi cá nhân đăng ký Use Case AI, cập nhật tiến độ; '
              'hội đồng và teamlead chấm điểm; hệ thống tổng hợp KPI và xếp hạng.')
add_body(doc, 'Tài liệu này gồm 2 phần theo vai trò:', bold=True)
add_bullets(doc, [
    'NHÂN SỰ (thành viên): đăng ký Use Case, cập nhật tuần, dùng Thư viện AI & xác nhận tái dùng, xem điểm.',
    'TEAMLEAD: chấm điểm US (hội đồng), chấm điểm cá nhân theo tháng, đọc KPI tổng hợp & Heatmap.',
])

# ── 2. Đăng nhập ──
add_heading(doc, '2. Đăng nhập hệ thống', 1)
add_step_box(doc, 1, 'Mở nền tảng & đăng nhập', 'Dùng chung tài khoản với SHTD Dashboard (username + mật khẩu). '
             'Nếu quên/đổi mật khẩu, dùng mục "Đổi mật khẩu" trên thanh menu.')
add_image(doc, '01_login.png', 'Hình 1 — Màn hình đăng nhập')
add_note_box(doc, 'Vai trò (nhân sự / teamlead / admin) được xác định tự động theo tài khoản. Menu bên trái hiển thị '
                  'đúng chức năng theo vai trò của bạn.', 'info')

# ══════════ PHẦN NHÂN SỰ ══════════
add_heading(doc, '3. Dành cho NHÂN SỰ', 1)
add_image(doc, '02_home_member.png', 'Hình 2 — Trang chủ (vai trò nhân sự)')

add_heading(doc, '3.1  Đăng ký Use Case AI', 2)
add_step_box(doc, 1, 'Chọn Workflow → Use Case', 'Vào "Đăng ký Use Case". Bước 1 chọn Workflow (lọc theo Team của bạn '
             '+ Workflow chung), rồi chọn Use Case tương ứng; hoặc chọn "Khác — nhập tự do" nếu chưa có trong danh mục.')
add_step_box(doc, 2, 'Điền 4 bước', 'Nghiệp vụ → AI & Prompt → Demo (link minh chứng) → Hướng dẫn. Bấm Gửi để nộp use case.')
add_image(doc, '03_register_workflow.png', 'Hình 3 — Bước 1: chọn Workflow và Use Case theo danh mục chuẩn')

add_heading(doc, '3.2  Cập nhật tiến độ tuần', 2)
add_step_box(doc, 1, 'Cập nhật kết quả áp dụng', 'Vào "Cập nhật tuần", chọn use case của bạn và điền tiến độ / kết quả '
             'ứng dụng trong tuần. Đây là căn cứ đo mức độ áp dụng AI thực tế.')
add_image(doc, '04_weekly_update.png', 'Hình 4 — Cập nhật tiến độ tuần')

add_heading(doc, '3.3  Thư viện AI & xác nhận tái dùng', 2)
add_body(doc, 'Thư viện gom các use case đã duyệt kèm prompt/workflow để mọi người tham khảo và tái dùng.')
add_step_box(doc, 1, 'Tham khảo & sao chép prompt', 'Mở use case trong Thư viện, xem cách làm và bấm Copy Prompt để dùng lại.')
add_step_box(doc, 2, 'Xác nhận đã tái dùng', 'Nếu bạn áp dụng lại một use case của người khác, bấm "Tôi đã tái dùng". '
             'Khi ≥ 3 người xác nhận, chủ use case được ghi nhận điểm "lan tỏa" (M-KPI-4).')
add_image(doc, '05_library_reuse.png', 'Hình 5 — Thư viện AI và xác nhận tái dùng')
add_note_box(doc, 'Không thể tự xác nhận tái dùng use case của chính mình.', 'warn')

add_heading(doc, '3.4  Xem điểm & Leaderboard', 2)
add_body(doc, 'Vào "Leaderboard" để xem xếp hạng điểm US (hội đồng), điểm cá nhân và KPI tổng hợp. '
              'Điểm cá nhân của bạn do teamlead chấm theo từng tháng (xem Phần 4).')
add_image(doc, '10_leaderboard_kpi.png', 'Hình 6 — Leaderboard: tab KPI tổng hợp')

add_heading(doc, '3.5  Đổi mật khẩu', 2)
add_image(doc, '12_change_password.png', 'Hình 7 — Đổi mật khẩu tự phục vụ')

# ══════════ PHẦN TEAMLEAD ══════════
doc.add_page_break()
add_heading(doc, '4. Dành cho TEAMLEAD', 1)
add_body(doc, 'Ngoài các chức năng của nhân sự, teamlead có thêm: chấm điểm US (hội đồng), chấm điểm cá nhân theo tháng, '
              'và đọc KPI tổng hợp / Heatmap của team.')

add_heading(doc, '4.1  Chấm điểm US — Hội đồng', 2)
add_step_box(doc, 1, 'Mở Hàng đợi Review', 'Vào "Hàng đợi Review" — danh sách use case chờ hội đồng chấm, kèm tiến độ n/4 '
             '(số thành viên hội đồng đã chấm).')
add_image(doc, '08_review_list.png', 'Hình 8 — Hàng đợi Review')
add_step_box(doc, 2, 'Chấm 3 tiêu chí (0–10)', 'Mở 1 use case, kéo 3 thanh: Tiết kiệm thời gian (30%) · Tự động hóa (40%) · '
             'Sáng tạo (30%). Giá trị hiển thị ngay trên thanh. Xem dòng "Bằng chứng (EVD)" để đối chiếu minh chứng, rồi bấm Gửi.')
add_image(doc, '09_review_panel.png', 'Hình 9 — Bảng chấm điểm US hội đồng (3 tiêu chí + dòng EVD)')
add_note_box(doc, 'Điểm US cuối cùng của use case = trung bình điểm của các thành viên hội đồng đã chấm.', 'info')

add_heading(doc, '4.2  Chấm điểm cá nhân theo THÁNG', 2)
add_body(doc, 'Điểm năng lực (M-KPI-2) được chấm theo TỪNG THÁNG. Điểm cuối kỳ = trung bình các tháng đã chấm.')
add_step_box(doc, 1, 'Mở "Chấm điểm cá nhân"', 'Danh sách thành viên team kèm điểm CN (trung bình các tháng) và số tháng đã chấm. '
             'Bấm "Chấm/Sửa" ở người cần chấm.')
add_image(doc, '06_personal_list.png', 'Hình 10 — Danh sách chấm điểm cá nhân')
add_step_box(doc, 2, 'Chọn kỳ (tháng) & chấm 4 tiêu chí', 'Trong bảng chấm, chọn "Kỳ chấm (tháng)". Tháng chưa chấm mặc định 0; '
             'tháng đã chấm hiển thị điểm cũ để sửa. Kéo 4 thanh: Đa dạng 30% · Thành thạo AI 20% · Chất lượng SP 30% · Số lượng đủ 20%.')
add_step_box(doc, 3, 'Xem các nhóm điểm & lưu', 'Bảng hiển thị rõ: M-KPI-1 (điểm US do hội đồng chấm — chỉ xem) · M-KPI-2 (bạn chấm) · '
             'KPI khác (khóa học/lan tỏa/điểm trừ) · KPI tổng hợp dự kiến · dòng EVD. Bấm "Lưu điểm tháng này".')
add_image(doc, '07_personal_panel.png', 'Hình 11 — Bảng chấm điểm cá nhân theo tháng (nhóm điểm rõ + thanh kéo hiển thị giá trị + EVD)')
add_note_box(doc, 'Khóa học / lan tỏa / số milestone chậm chỉ cần nhập 1 lần (hệ thống lấy giá trị mới nhất). '
                  'Riêng 4 tiêu chí năng lực được chấm & lưu theo từng tháng.', 'tip')

add_heading(doc, '4.3  Đọc KPI tổng hợp & Heatmap', 2)
add_body(doc, 'Trên Leaderboard, tab "KPI tổng hợp" và "KPI Teamlead" thể hiện điểm cuối; tab "Heatmap" hiển thị '
              'lưới team × cá nhân theo màu (xanh ≥85 · xanh lá ≥70 · vàng 50–69 · đỏ <50).')
add_image(doc, '11_leaderboard_heatmap.png', 'Hình 12 — Heatmap KPI team & cá nhân (kèm thẻ KPI PM)')

# ── 5. Bộ điểm & công thức ──
doc.add_page_break()
add_heading(doc, '5. Bộ điểm & công thức KPI', 1)
add_body(doc, 'KPI thành viên (Member) tổng hợp từ 4 cấu phần, trừ điểm milestone chậm:', bold=True)
add_table(doc,
    ['Cấu phần', 'Trọng số', 'Người chấm', 'Ghi chú'],
    [
        ['M-KPI-1 — Điểm US', '40%', 'Hội đồng (4 teamlead)', 'Bình quân điểm hội đồng các UC do cá nhân sở hữu'],
        ['M-KPI-2 — Năng lực', '30%', 'Teamlead (theo tháng)', 'Trung bình các tháng đã chấm (4 tiêu chí)'],
        ['M-KPI-3 — Khóa học', '15%', 'Teamlead', 'Mỗi khóa 25%, khóa trả phí ×2 (tối đa 100%)'],
        ['M-KPI-4 — Lan tỏa', '15%', 'Teamlead / tự động', 'Đạt = có buổi chia sẻ HOẶC ≥3 người tái dùng UC'],
        ['Điểm trừ — Milestone chậm', '−2%/mốc', 'Teamlead', 'Tối đa −10%'],
    ])
add_body(doc, 'Ngưỡng xếp hạng (thang 100):', bold=True)
add_table(doc,
    ['Xếp hạng', 'Điểm', 'Màu'],
    [
        ['Xuất sắc (Top)', '≥ 85', 'Tím'],
        ['Tốt (Strong)', '≥ 70', 'Xanh lá'],
        ['Trung bình (Average)', '≥ 50', 'Vàng'],
        ['Cần cải thiện', '< 50', 'Đỏ'],
    ])
add_note_box(doc, 'Teamlead KPI = 60% KPI cá nhân của teamlead + 40% tỷ lệ thành viên team đạt ≥ 70 điểm.', 'info')

# ── 6. FAQ ──
add_heading(doc, '6. Câu hỏi thường gặp', 1)
faqs = [
    ('Không thấy Workflow trong danh mục khi đăng ký?',
     'Chọn "Khác — nhập tự do" để nhập Use Case mới; hoặc báo admin bổ sung Workflow vào danh mục.'),
    ('Điểm cá nhân của tôi thấp vì mới chấm 1 tháng?',
     'Điểm cuối kỳ là trung bình các tháng ĐÃ chấm — tháng chưa chấm không bị tính 0. Điểm sẽ phản ánh dần theo các tháng.'),
    ('Teamlead chấm nhầm tháng thì sao?',
     'Mở lại người đó, chọn đúng tháng và chấm lại — hệ thống ghi đè điểm của tháng đó.'),
    ('Dòng EVD để làm gì?',
     'Là link minh chứng (ổ share) để đối chiếu khi chấm; hiện chỉ hiển thị, không chỉnh sửa tại màn chấm.'),
]
for q, a in faqs:
    p = doc.add_paragraph(); r = p.add_run('• ' + q); r.font.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = PURPLE
    add_body(doc, '   ' + a, size=10, color=DARK)

doc.save(OUT)
print('WROTE', OUT)
