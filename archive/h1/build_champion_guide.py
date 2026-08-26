"""
Tạo hướng dẫn sử dụng dành cho Champion — AI-US-SPTD v3.10.3
Output: HDSD_Champion_AI_USSPTD.docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

SCREENSHOTS = os.path.join(os.path.dirname(__file__), 'screenshots')
OUT_FILE = os.path.join(os.path.dirname(__file__), 'HDSD_Champion_AI_USSPTD.docx')

PURPLE = RGBColor(0x6B, 0x21, 0xA8)
DARK   = RGBColor(0x1F, 0x25, 0x37)
GRAY   = RGBColor(0x6B, 0x72, 0x80)
GREEN  = RGBColor(0x05, 0x96, 0x69)
AMBER  = RGBColor(0xD9, 0x77, 0x06)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p


def add_body(doc, text, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_image(doc, filename, caption, width=Cm(15)):
    path = os.path.join(SCREENSHOTS, filename)
    if not os.path.exists(path):
        doc.add_paragraph(f'[Hình: {caption} — file không tồn tại: {filename}]')
        return
    doc.add_picture(path, width=width)
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.runs[0] if cap.runs else cap.add_run(caption)
    cap_run.font.size = Pt(9)
    cap_run.font.italic = True
    cap_run.font.color.rgb = GRAY


def add_step_box(doc, step_num, title, desc):
    """Tạo bảng 1 hàng để hiển thị bước hướng dẫn nổi bật."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # số bước
    c0 = tbl.rows[0].cells[0]
    c0.width = Cm(1.5)
    set_cell_bg(c0, '6B21A8')
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(str(step_num))
    r0.font.size = Pt(16)
    r0.font.bold = True
    r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # nội dung
    c1 = tbl.rows[0].cells[1]
    p1 = c1.paragraphs[0]
    r1a = p1.add_run(title + '\n')
    r1a.font.bold = True
    r1a.font.size = Pt(11)
    r1b = p1.add_run(desc)
    r1b.font.size = Pt(10)
    r1b.font.color.rgb = GRAY
    doc.add_paragraph()


def add_note_box(doc, text, note_type='info'):
    """Khung ghi chú / cảnh báo."""
    color_map = {'info': 'EDE9FE', 'warn': 'FEF3C7', 'tip': 'D1FAE5'}
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, color_map.get(note_type, 'EDE9FE'))
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    doc.add_paragraph()


def add_score_table(doc):
    headers = ['Tiêu chí', 'Điểm tối đa', 'Người chấm', 'Mô tả']
    rows = [
        ('Hiệu quả (Efficiency)', '20 pt', 'Tự động', 'Giảm thời gian / công sức thực hiện'),
        ('Độ phổ biến (Adoption)', '20 pt', 'Tự động', 'Mức độ sử dụng trong team'),
        ('Khả năng tái dùng (Reuse)', '20 pt', 'Tự động', 'Có thể áp dụng cho team/bộ phận khác'),
        ('Tần suất sử dụng (Frequency)', '15 pt', 'Tự động', 'Sử dụng bao nhiêu lần / tuần'),
        ('Tài liệu hóa (Docs)', '5 pt', 'Tự động', 'Có hướng dẫn, prompt mẫu rõ ràng'),
        ('Chất lượng (Quality)', '10 pt', '★ Champion', 'Mức độ hoàn thiện, độ chính xác AI'),
        ('Giá trị kinh doanh (Business Value)', '10 pt', '★ Champion', 'Tác động thực tế đến KPI, doanh thu'),
        ('Sáng tạo (Innovation)', '10 pt', '★ Champion', 'Ý tưởng mới, cách tiếp cận độc đáo'),
    ]
    tbl = doc.add_table(rows=1 + len(rows), cols=4)
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, '6B21A8')
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
    for ri, row_data in enumerate(rows, start=1):
        for ci, val in enumerate(row_data):
            cell = tbl.rows[ri].cells[ci]
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            if '★' in val:
                run.font.bold = True
                run.font.color.rgb = PURPLE
            if ci == 1 and '★' in rows[ri-1][2]:
                set_cell_bg(cell, 'EDE9FE')
    doc.add_paragraph()


def add_rank_table(doc):
    headers = ['Xếp hạng', 'Điểm', 'Màu sắc', 'Ý nghĩa']
    rows = [
        ('TOP PERFORMER',       '≥ 85 pt', 'Tím (Purple)',  'Xuất sắc — Tiếp tục phát huy'),
        ('STRONG CONTRIBUTOR',  '≥ 70 pt', 'Xanh (Green)',  'Tốt — Duy trì và mở rộng'),
        ('AVERAGE',             '≥ 50 pt', 'Vàng (Amber)',  'Trung bình — Cần cải thiện chất lượng'),
        ('CẦN CẢI THIỆN',       '< 50 pt', 'Đỏ (Red)',      'Cần xem xét lại và nâng cấp use case'),
    ]
    tbl = doc.add_table(rows=1 + len(rows), cols=4)
    tbl.style = 'Table Grid'
    header_colors = ['6B21A8', '6B21A8', '6B21A8', '6B21A8']
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, header_colors[i])
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
    row_colors = ['EDE9FE', 'D1FAE5', 'FEF3C7', 'FEE2E2']
    for ri, (rdata, rc) in enumerate(zip(rows, row_colors), start=1):
        for ci, val in enumerate(rdata):
            cell = tbl.rows[ri].cells[ci]
            set_cell_bg(cell, rc)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            if ci == 0:
                run.font.bold = True
    doc.add_paragraph()


# ─────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

# ── TRANG BÌA ──────────────────────────────
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run('\n\n\n🚀 HƯỚNG DẪN SỬ DỤNG\nHỆ THỐNG AI USE CASE\ndành cho CHAMPION')
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = PURPLE

doc.add_paragraph()
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run('Trung tâm Sản phẩm & Dịch vụ — TT SPTD\nPhiên bản: v3.10.3  |  Ngày ban hành: ' + datetime.date.today().strftime('%d/%m/%Y'))
r2.font.size = Pt(12)
r2.font.color.rgb = GRAY

doc.add_page_break()

# ── 1. GIỚI THIỆU ──────────────────────────
add_heading(doc, '1. Giới thiệu', level=1, color=PURPLE)
add_body(doc,
    'Hệ thống AI Use Case (AI-US-SPTD) là nền tảng quản lý và đánh giá các ứng dụng AI '
    'trong nội bộ Trung tâm SP&GPTD. Hệ thống vừa được cập nhật với vai trò mới: '
    'Champion — người đại diện team chịu trách nhiệm thẩm định chất lượng use case AI của team mình.')

add_body(doc, 'Vai trò Champion có quyền hạn gì?', bold=True)
for item in [
    '✅  Xem và duyệt tất cả use case do thành viên team nộp lên',
    '✅  Chấm điểm chất lượng (30 điểm thủ công): Quality · Business Value · Innovation',
    '✅  Theo dõi bảng xếp hạng (Leaderboard) toàn hệ thống',
    '✅  Đăng ký use case của bản thân',
    '❌  Không thể xem use case của team khác trong hàng đợi duyệt',
    '❌  Không có quyền quản trị người dùng (chỉ Admin)',
]:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.size = Pt(10.5)

doc.add_paragraph()
add_note_box(doc, '💡 Lưu ý: Tài khoản Champion do Admin tạo và cấp quyền. '
             'Nếu bạn chưa đăng nhập được với role Champion, hãy liên hệ quản trị viên hệ thống.', 'info')

# ── 2. ĐĂNG NHẬP ──────────────────────────
add_heading(doc, '2. Đăng nhập hệ thống', level=1, color=PURPLE)
add_body(doc, 'Truy cập hệ thống tại địa chỉ:')
p_url = doc.add_paragraph()
r_url = p_url.add_run('https://tuanttstb-debug.github.io/ai-usecase-platform/')
r_url.font.bold = True
r_url.font.color.rgb = PURPLE
r_url.font.size = Pt(11)

doc.add_paragraph()
add_step_box(doc, 1, 'Nhập tên đăng nhập',
             'Dùng tên đăng nhập nội bộ TT SPTD (ví dụ: tuantt4). Không nhập mật khẩu — hệ thống xác thực qua danh sách nội bộ.')
add_step_box(doc, 2, 'Nhấn nút "Đăng nhập →"',
             'Nếu tài khoản hợp lệ và đã được cấp quyền Champion, bạn sẽ được chuyển vào Trang chủ.')

add_image(doc, '01_login.png', 'Hình 1 — Màn hình đăng nhập hệ thống AI-US-SPTD')

add_note_box(doc, '⚠️  Nếu bạn thấy thông báo lỗi hoặc bị đưa về trang đăng nhập sau khi nhập tên, '
             'hãy kiểm tra lại chính xác tên đăng nhập (phân biệt hoa/thường) và liên hệ Admin.', 'warn')

# ── 3. GIAO DIỆN CHAMPION ──────────────────
add_heading(doc, '3. Giao diện dành cho Champion', level=1, color=PURPLE)
add_body(doc,
    'Sau khi đăng nhập thành công với tài khoản Champion, bạn sẽ thấy giao diện trang chủ '
    'với thanh điều hướng bên trái. Góc trên bên phải và footer sidebar sẽ hiển thị tên và role "Champion" của bạn.')

add_image(doc, '02_home_champion.png', 'Hình 2 — Trang chủ với vai trò Champion (sidebar hiển thị "Hàng đợi Review")')

add_body(doc, 'Các mục điều hướng Champion có thể truy cập:', bold=True)
nav_items = [
    ('📋  Đăng ký Use Case', 'Nộp use case AI mới của bản thân'),
    ('📁  Use Case của tôi', 'Xem danh sách và trạng thái các UC bạn đã nộp'),
    ('🔍  Hàng đợi Review', '★ CHỨC NĂNG CHÍNH — Duyệt và chấm điểm UC của team'),
    ('🏆  Leaderboard', 'Bảng xếp hạng hiệu suất AI toàn trung tâm'),
    ('📰  Cập nhật tuần', 'Bản tin nội bộ hàng tuần về AI'),
]
tbl = doc.add_table(rows=len(nav_items), cols=2)
tbl.style = 'Table Grid'
for i, (name, desc) in enumerate(nav_items):
    c0 = tbl.rows[i].cells[0]
    c1 = tbl.rows[i].cells[1]
    if i % 2 == 0:
        set_cell_bg(c0, 'F5F3FF')
        set_cell_bg(c1, 'F5F3FF')
    r0 = c0.paragraphs[0].add_run(name)
    r0.font.bold = True
    r0.font.size = Pt(10)
    r1 = c1.paragraphs[0].add_run(desc)
    r1.font.size = Pt(10)
    if '★' in name:
        r0.font.color.rgb = PURPLE
doc.add_paragraph()

# ── 4. NHIỆM VỤ CHÍNH: REVIEW USE CASE ────
add_heading(doc, '4. Nhiệm vụ chính: Review Use Case của team', level=1, color=PURPLE)
add_body(doc,
    'Đây là nhiệm vụ quan trọng nhất của Champion. Mỗi khi thành viên trong team nộp use case mới, '
    'use case đó sẽ xuất hiện trong "Hàng đợi Review" và chờ Champion chấm điểm.')

add_heading(doc, '4.1  Truy cập Hàng đợi Review', level=2)
add_body(doc, 'Nhấn vào "Hàng đợi Review" trên thanh điều hướng trái. '
         'Hệ thống chỉ hiển thị use case của team bạn — không lẫn use case team khác.')

add_image(doc, '05b_review_queue_admin.png',
          'Hình 3 — Hàng đợi Review: danh sách UC chờ đánh giá, lọc theo trạng thái và team')

add_body(doc, 'Các tab lọc trạng thái:', bold=True)
for item in [
    '🟡  Chờ đánh giá — UC chưa có điểm thủ công, cần Champion chấm',
    '🔵  Đang review — Champion đang trong quá trình đánh giá',
    '🟢  Đã hoàn thành — UC đã được chấm điểm đầy đủ',
]:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.size = Pt(10.5)
doc.add_paragraph()

add_heading(doc, '4.2  Quy trình chấm điểm', level=2)
add_step_box(doc, 1, 'Nhấn nút "Review" bên phải của use case cần đánh giá',
             'Panel chi tiết sẽ trượt ra từ bên phải màn hình.')
add_step_box(doc, 2, 'Đọc kỹ thông tin use case',
             'Xem tên UC, mô tả bài toán, công cụ AI sử dụng, prompt mẫu và kết quả demo (nếu có).')
add_step_box(doc, 3, 'Chấm điểm 3 tiêu chí (tổng 30 điểm)',
             'Kéo thanh slider cho từng tiêu chí: Quality (10đ) · Business Value (10đ) · Innovation (10đ).')
add_step_box(doc, 4, 'Nhấn "Lưu đánh giá"',
             'Điểm thủ công sẽ được cộng vào điểm tự động (70đ) để tính tổng điểm 100đ. UC chuyển trạng thái "Đã hoàn thành".')

add_note_box(doc,
    '💡 Mẹo đánh giá: Hãy mở thử use case hoặc kiểm tra kết quả thực tế trước khi chấm điểm. '
    'Điểm chất lượng (Quality) nên phản ánh mức độ AI hoạt động đúng và ổn định, '
    'không chỉ dựa vào mô tả văn bản.', 'tip')

# ── 5. THEO DÕI USE CASE CỦA BẠN ──────────
add_heading(doc, '5. Theo dõi Use Case của bản thân', level=1, color=PURPLE)
add_body(doc,
    'Ngoài nhiệm vụ duyệt, Champion cũng có thể nộp use case cá nhân như mọi thành viên khác. '
    'Vào "Use Case của tôi" để xem danh sách và trạng thái các UC bạn đã đăng ký.')

add_image(doc, '03_dashboard_my.png',
          'Hình 4 — Màn hình "Use Case của tôi": danh sách UC cá nhân, trạng thái và ngày nộp')

add_body(doc, 'Tab "Khám phá" cho phép xem toàn bộ use case của tất cả team trong tổ chức, '
         'kèm điểm số và xếp hạng — rất hữu ích để tham khảo và học hỏi.')
add_image(doc, '04_dashboard_explore.png',
          'Hình 5 — Tab "Khám phá": xem toàn bộ UC của tổ chức kèm điểm và team')

# ── 6. HỆ THỐNG ĐIỂM & XẾP HẠNG ──────────
add_heading(doc, '6. Hệ thống điểm và xếp hạng', level=1, color=PURPLE)
add_body(doc,
    'Mỗi use case được chấm tổng cộng 100 điểm, gồm 70 điểm tự động (hệ thống tính) '
    'và 30 điểm thủ công do Champion chấm.')

add_heading(doc, '6.1  Bảng tiêu chí chấm điểm (100 điểm)', level=2)
add_score_table(doc)

add_heading(doc, '6.2  Ngưỡng xếp hạng', level=2)
add_rank_table(doc)

add_image(doc, '09_leaderboard.png',
          'Hình 6 — Leaderboard: bảng xếp hạng Top Performers toàn trung tâm theo điểm số')

add_note_box(doc,
    '⚠️  Lưu ý quan trọng: Điểm thủ công của Champion có tác động lớn đến xếp hạng cuối cùng. '
    'Một use case 40đ tự động + 30đ Champion = 70đ (STRONG CONTRIBUTOR). '
    'Hãy đánh giá khách quan và nhất quán giữa các thành viên.', 'warn')

# ── 7. ĐĂNG KÝ USE CASE MỚI ───────────────
add_heading(doc, '7. Đăng ký Use Case mới', level=1, color=PURPLE)
add_body(doc,
    'Champion cũng có thể đăng ký use case AI của bản thân. '
    'Form đăng ký gồm 4 bước, mỗi bước cần điền đầy đủ thông tin trước khi chuyển bước tiếp theo.')

add_image(doc, '07_register_form.png',
          'Hình 7 — Form đăng ký Use Case: wizard 4 bước (Nghiệp vụ → AI & Prompt → Demo → Hướng dẫn)')

add_body(doc, 'Mô tả 4 bước đăng ký:', bold=True)
steps = [
    ('Bước 1 — Nghiệp vụ', 'Tên Use Case, người đăng ký, team, lĩnh vực, bài toán cần giải quyết, kết quả kỳ vọng'),
    ('Bước 2 — AI & Prompt', 'Công cụ AI sử dụng, prompt mẫu, cách thức tích hợp vào quy trình làm việc'),
    ('Bước 3 — Demo', 'Ảnh chụp màn hình hoặc mô tả kết quả thực tế đã đạt được'),
    ('Bước 4 — Hướng dẫn', 'Hướng dẫn sử dụng ngắn gọn để thành viên khác có thể tái sử dụng'),
]
for step, desc in steps:
    p = doc.add_paragraph()
    r1 = p.add_run(f'{step}: ')
    r1.font.bold = True
    r1.font.color.rgb = PURPLE
    r1.font.size = Pt(10.5)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)
doc.add_paragraph()

add_note_box(doc,
    '💡 Gợi ý: Điền đầy đủ và chi tiết phần "Prompt mẫu" và "Hướng dẫn" sẽ giúp use case '
    'đạt điểm cao hơn ở tiêu chí Tài liệu hóa (5đt tự động) và '
    'được Champion đánh giá Quality tốt hơn.', 'tip')

# ── 8. CÂU HỎI THƯỜNG GẶP ─────────────────
add_heading(doc, '8. Câu hỏi thường gặp (FAQ)', level=1, color=PURPLE)
faqs = [
    ('Tôi không thấy use case của team trong Hàng đợi Review?',
     'Kiểm tra lại tên team trong tài khoản của bạn có khớp chính xác với tên team của người nộp UC không (phân biệt hoa/thường). Liên hệ Admin để kiểm tra cấu hình team.'),
    ('Tôi đã chấm điểm nhưng điểm không thay đổi?',
     'Sau khi nhấn "Lưu đánh giá", hãy làm mới trang (F5) và kiểm tra lại. Nếu vẫn không cập nhật, có thể hệ thống đang trong quá trình đồng bộ — thử lại sau 1-2 phút.'),
    ('Tôi có thể sửa điểm đã chấm không?',
     'Hiện tại chỉ Admin mới có thể chỉnh sửa điểm đã lưu. Hãy kiểm tra kỹ trước khi nhấn Lưu. Nếu cần điều chỉnh, liên hệ Admin.'),
    ('Use case "Chờ đánh giá" đã lâu nhưng tôi không thấy?',
     'Hàng đợi Review chỉ hiển thị UC của team bạn. Nếu không thấy, có thể team của người nộp khác team bạn đang quản lý.'),
    ('Tôi bị đăng xuất sau một thời gian?',
     'Phiên đăng nhập có thời hạn. Đăng nhập lại bằng tên đăng nhập của bạn là đủ.'),
]
for q, a in faqs:
    p_q = doc.add_paragraph()
    rq = p_q.add_run(f'❓  {q}')
    rq.font.bold = True
    rq.font.color.rgb = DARK
    rq.font.size = Pt(10.5)
    p_a = doc.add_paragraph()
    ra = p_a.add_run(f'    → {a}')
    ra.font.size = Pt(10.5)
    ra.font.color.rgb = GRAY
    doc.add_paragraph()

# ── FOOTER ─────────────────────────────────
doc.add_page_break()
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = p_foot.add_run(
    'Hệ thống AI Use Case — TT SPTD\n'
    'Mọi thắc mắc liên hệ: Admin hệ thống hoặc kênh nội bộ TT SPTD\n'
    f'Tài liệu phát hành: {datetime.date.today().strftime("%d/%m/%Y")}  |  Phiên bản: v3.10.3'
)
rf.font.size = Pt(9)
rf.font.color.rgb = GRAY

doc.save(OUT_FILE)
print(f'[OK] Da tao: {OUT_FILE}')
